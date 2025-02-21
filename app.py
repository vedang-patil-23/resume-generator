import streamlit as st
from docx import Document
import base64
import os
import uuid

def generate_resume(data):
    doc = Document()
    doc.add_heading(data["name"], 0)
    doc.add_paragraph(f"Phone: {data['phone']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Introduction: {data['intro']}")

    doc.add_heading("Education", level=1)
    for edu in data["education"]:
        doc.add_paragraph(f"{edu['degree']} - {edu['university']} ({edu['years']})")

    doc.add_heading("Experience", level=1)
    for exp in data["experience"]:
        doc.add_paragraph(f"{exp['title']} at {exp['company']} ({exp['years']})")
        doc.add_paragraph(exp["description"])

    doc.add_heading("Projects", level=1)
    for proj in data["projects"]:
        doc.add_paragraph(f"{proj['name']}: {proj['description']}")

    doc.add_heading("Publications", level=1)
    for pub in data["publications"]:
        doc.add_paragraph(f"{pub['title']} - {pub['reference']}")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data["skills"]))

    doc.add_heading("Co-Curricular Activities", level=1)
    for act in data["coCurricular"]:
        doc.add_paragraph(f"{act['activity']} ({act['year']}), {act['position']}")
        doc.add_paragraph(act["description"])

    file_path = f"Resume_{uuid.uuid4().hex}.docx"
    doc.save(file_path)
    return file_path

st.title("📄 Resume Generator")

with st.form("resume_form"):
    name = st.text_input("Full Name *", "")
    phone = st.text_input("Phone Number *", "")
    email = st.text_input("Email *", "")
    intro = st.text_area("Introduction *", "")

    # Education
    education = []
    if st.checkbox("Add Education"):
        edu_degree = st.text_input("Degree")
        edu_university = st.text_input("University")
        edu_years = st.text_input("Years")
        education.append({"degree": edu_degree, "university": edu_university, "years": edu_years})

    # Experience
    experience = []
    if st.checkbox("Add Experience"):
        exp_title = st.text_input("Job Title")
        exp_company = st.text_input("Company")
        exp_years = st.text_input("Years Worked")
        exp_desc = st.text_area("Job Description")
        experience.append({"title": exp_title, "company": exp_company, "years": exp_years, "description": exp_desc})

    # Projects
    projects = []
    if st.checkbox("Add Projects"):
        proj_name = st.text_input("Project Name")
        proj_desc = st.text_area("Project Description")
        projects.append({"name": proj_name, "description": proj_desc})

    # Publications
    publications = []
    if st.checkbox("Add Publications"):
        pub_title = st.text_input("Publication Title")
        pub_ref = st.text_input("Reference")
        publications.append({"title": pub_title, "reference": pub_ref})

    # Skills
    skills = st.text_area("Skills (comma separated)").split(",")

    # Co-Curricular Activities
    coCurricular = []
    if st.checkbox("Add Co-Curricular Activities"):
        co_activity = st.text_input("Activity")
        co_year = st.text_input("Year")
        co_position = st.text_input("Position")
        co_desc = st.text_area("Description")
        coCurricular.append({"activity": co_activity, "year": co_year, "position": co_position, "description": co_desc})

    submitted = st.form_submit_button("Generate Resume")

if submitted:
    form_data = {
        "name": name,
        "phone": phone,
        "email": email,
        "intro": intro,
        "education": education,
        "experience": experience,
        "projects": projects,
        "publications": publications,
        "skills": skills,
        "coCurricular": coCurricular,
    }

    resume_file = generate_resume(form_data)

    with open(resume_file, "rb") as file:
        docx_data = file.read()
        b64 = base64.b64encode(docx_data).decode()
        st.markdown(f'<a href="data:file/docx;base64,{b64}" download="Resume.docx">📥 Download Resume</a>', unsafe_allow_html=True)

    os.remove(resume_file)  # Clean up file after download
